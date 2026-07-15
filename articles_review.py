from __future__ import annotations

import io
import json
import re
import struct
import zlib
from datetime import datetime
from pathlib import Path
from typing import Any

import streamlit as st
from pypdf import PdfReader

from customer_history import save_customer_event
from utils import ROOT_DIR, get_user_dirs

CHECKLIST_PATH = ROOT_DIR / "data" / "articles_review_checklist.json"


def _normalize_business_no(value: Any) -> str:
    return re.sub(r"[^0-9]", "", str(value or ""))


def _load_json(path: Path, default: Any) -> Any:
    if not path.exists():
        return default
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return default


def _save_json(path: Path, value: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(value, ensure_ascii=False, indent=2, default=str),
        encoding="utf-8",
    )


def _review_path(user_id: str) -> Path:
    return get_user_dirs(user_id)["base"] / "articles_reviews.json"


def _extract_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _extract_docx(data: bytes) -> str:
    from docx import Document
    document = Document(io.BytesIO(data))
    blocks = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            blocks.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(blocks)


def _extract_hwp(data: bytes) -> str:
    import olefile

    ole = olefile.OleFileIO(io.BytesIO(data))
    file_header = ole.openstream("FileHeader").read()
    compressed = bool(file_header[36] & 1)

    sections = sorted(
        path for path in ole.listdir()
        if len(path) == 2 and path[0] == "BodyText"
        and path[1].startswith("Section")
    )
    texts: list[str] = []
    for section in sections:
        raw = ole.openstream(section).read()
        if compressed:
            raw = zlib.decompress(raw, -15)
        offset = 0
        while offset + 4 <= len(raw):
            header = struct.unpack_from("<I", raw, offset)[0]
            tag_id = header & 0x3FF
            size = (header >> 20) & 0xFFF
            offset += 4
            if size == 0xFFF:
                if offset + 4 > len(raw):
                    break
                size = struct.unpack_from("<I", raw, offset)[0]
                offset += 4
            payload = raw[offset:offset + size]
            offset += size
            if tag_id == 67:
                try:
                    texts.append(payload.decode("utf-16le", errors="ignore"))
                except Exception:
                    continue
    return "\n".join(texts)


def extract_articles_text(filename: str, data: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(data)
    if suffix == ".docx":
        return _extract_docx(data)
    if suffix == ".hwp":
        return _extract_hwp(data)
    if suffix in {".txt", ".md"}:
        return data.decode("utf-8", errors="ignore")
    raise ValueError("지원 형식은 PDF, HWP, DOCX, TXT입니다.")


def analyze_articles(text: str) -> dict[str, Any]:
    checklist = _load_json(CHECKLIST_PATH, {}).get("items", [])
    normalized = re.sub(r"\s+", "", text).lower()
    items = []
    total_weight = 0
    earned = 0

    for item in checklist:
        weight = int(item.get("weight", 0) or 0)
        total_weight += weight
        matched = [
            term for term in item.get("terms", [])
            if re.sub(r"\s+", "", str(term)).lower() in normalized
        ]
        ratio = min(len(matched) / max(min(len(item.get("terms", [])), 2), 1), 1)
        status = "충분" if ratio >= 1 else "부분반영" if matched else "미확인"
        score = round(weight * ratio)
        earned += score
        items.append({
            "id": item.get("id"),
            "title": item.get("title"),
            "category": item.get("category"),
            "weight": weight,
            "score": score,
            "status": status,
            "matched_terms": matched,
            "recommendation": (
                "관련 조항과 별도 지급규정의 결의·시행일·적용대상을 함께 확인하세요."
                if status == "충분"
                else "정관 조항 및 별도 규정의 신설·보완 여부를 검토하세요."
            ),
        })

    score = round(earned / total_weight * 100) if total_weight else 0
    priorities = [
        item for item in items
        if item["status"] != "충분"
    ]
    priorities.sort(key=lambda row: row["weight"], reverse=True)

    consulting_map = {
        "executive_retirement": "임원퇴직금·CEO 퇴직재원 컨설팅",
        "survivor_compensation": "유족보상금·경영인정기보험 컨설팅",
        "executive_compensation": "임원보수체계 정비",
        "executive_bonus": "임원성과급·상여금 규정 정비",
        "treasury_stock": "자기주식·가업승계 컨설팅",
        "stock_options": "스톡옵션·핵심인재 보상설계",
        "preferred_shares": "투자유치·종류주식 설계",
        "share_transfer": "주식이동·가업승계 사전정비",
    }
    opportunities = [
        consulting_map[item["id"]]
        for item in priorities
        if item["id"] in consulting_map
    ]

    return {
        "score": score,
        "reviewed_at": datetime.now().isoformat(timespec="seconds"),
        "items": items,
        "priority_items": priorities[:8],
        "consulting_opportunities": opportunities[:6],
        "text_length": len(text),
    }


def save_articles_review(
    user_id: str,
    business_no: str,
    company_name: str,
    filename: str,
    analysis: dict[str, Any],
) -> dict[str, Any]:
    path = _review_path(user_id)
    all_data = _load_json(path, {})
    if not isinstance(all_data, dict):
        all_data = {}
    key = _normalize_business_no(business_no) or company_name
    record = {
        "company_name": company_name,
        "business_no": business_no,
        "filename": filename,
        **analysis,
    }
    all_data[key] = record
    _save_json(path, all_data)

    detail = (
        f"정관점수: {analysis.get('score', 0)}점\n"
        f"우선개선: "
        + ", ".join(
            item.get("title", "")
            for item in analysis.get("priority_items", [])[:5]
        )
    )
    save_customer_event(
        user_id=user_id,
        business_no=business_no,
        company_name=company_name,
        event_id=f"articles-review-{analysis.get('reviewed_at', '')}",
        event_title=f"{analysis.get('reviewed_at', '')[:10]} 정관검토 완료",
        event_detail=detail,
        occurred_at=analysis.get("reviewed_at", ""),
        source="articles_review",
    )
    return record


def get_latest_articles_review(
    user_id: str,
    business_no: str,
    company_name: str = "",
) -> dict[str, Any]:
    data = _load_json(_review_path(user_id), {})
    key = _normalize_business_no(business_no) or company_name
    value = data.get(key, {}) if isinstance(data, dict) else {}
    return value if isinstance(value, dict) else {}


def render_articles_review(
    user_id: str,
    business_no: str,
    company_name: str,
) -> None:
    st.markdown("#### 정관 AI 검토")
    st.caption(
        "정관의 핵심 조항을 분석해 절세·퇴직·유족보상·가업승계·"
        "투자유치 관점의 보완사항을 확인합니다."
    )

    uploaded = st.file_uploader(
        "정관 파일 업로드",
        type=["pdf", "hwp", "docx", "txt"],
        key=f"articles_upload_{business_no or company_name}",
    )
    if uploaded is not None and st.button(
        "정관 분석 실행",
        type="primary",
        use_container_width=True,
        key=f"articles_analyze_{business_no or company_name}",
    ):
        try:
            with st.spinner("정관 조항을 추출하고 분석하고 있습니다..."):
                text = extract_articles_text(uploaded.name, uploaded.getvalue())
                if len(re.sub(r"\s+", "", text)) < 100:
                    raise ValueError(
                        "추출된 정관 텍스트가 너무 짧습니다. "
                        "스캔 PDF는 텍스트 PDF로 변환 후 업로드해주세요."
                    )
                analysis = analyze_articles(text)
                save_articles_review(
                    user_id,
                    business_no,
                    company_name,
                    uploaded.name,
                    analysis,
                )
            st.success("정관 분석과 기업히스토리 저장을 완료했습니다.")
            st.rerun()
        except Exception as exc:
            st.error(f"정관 분석 실패: {exc}")

    review = get_latest_articles_review(user_id, business_no, company_name)
    if not review:
        st.info("저장된 정관검토 결과가 없습니다.")
        return

    c1, c2, c3 = st.columns(3)
    c1.metric("정관 종합점수", f"{review.get('score', 0)}점")
    c2.metric("확인 조항", f"{len(review.get('items', []))}개")
    c3.metric("우선개선", f"{len(review.get('priority_items', []))}개")
    st.caption(
        f"분석파일: {review.get('filename', '-')} · "
        f"분석일: {str(review.get('reviewed_at', ''))[:19]}"
    )

    rows = []
    for item in review.get("items", []):
        rows.append({
            "분야": item.get("category"),
            "검토조항": item.get("title"),
            "상태": item.get("status"),
            "점수": f"{item.get('score', 0)}/{item.get('weight', 0)}",
            "확인문구": ", ".join(item.get("matched_terms", [])),
            "검토의견": item.get("recommendation"),
        })
    st.dataframe(rows, hide_index=True, use_container_width=True)

    left, right = st.columns(2)
    with left:
        st.markdown("##### 우선 개정 검토")
        for item in review.get("priority_items", []):
            st.write(f"- **{item.get('title')}**: {item.get('recommendation')}")
    with right:
        st.markdown("##### 컨설팅 기회")
        opportunities = review.get("consulting_opportunities", [])
        if opportunities:
            for item in opportunities:
                st.write(f"- {item}")
        else:
            st.write("- 핵심 조항이 전반적으로 확인되었습니다.")
